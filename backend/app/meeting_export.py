from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from typing import Any

MAX_TRANSCRIPT_CHARACTERS = 2_000_000
MAX_EXPORT_BYTES = 20_000_000
CONTENT_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
    "json": "application/json; charset=utf-8",
}
EXTENSIONS = {"pdf": "pdf", "docx": "docx", "md": "md", "txt": "txt", "json": "json"}


def safe_export_filename(title: str, date_value: str | None, export_format: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "-", title.strip()).strip("-._") or "meeting"
    stem = stem[:80]
    date_part = (date_value or "")[:10]
    suffix = f"-{date_part}" if re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_part) else ""
    return f"{stem}{suffix}.{EXTENSIONS[export_format]}"


def format_timestamp(milliseconds: int | None) -> str:
    seconds = max(0, int(milliseconds or 0) // 1000)
    return f"{seconds // 60:02d}:{seconds % 60:02d}"


def _translated(original: Any, translated: Any, language: str) -> tuple[Any, bool]:
    if language != "english":
        return original, False
    if translated not in (None, ""):
        return translated, False
    return original, original not in (None, "")


def build_export_document(data: dict[str, Any], sections: dict[str, bool], language: str) -> dict[str, Any]:
    meeting = data["meeting"]
    document: dict[str, Any] = {
        "title": meeting.get("title") or "Untitled meeting",
        "export_language": language,
        "translation_fallback_sections": [],
    }
    fallback: set[str] = set()

    if sections.get("meeting_details"):
        document["meeting_details"] = {
            key: meeting.get(key)
            for key in (
                "title", "description", "scheduled_at", "duration_seconds", "status",
                "detected_language", "transcript_language",
            )
        }
    if sections.get("tags"):
        document["tags"] = data.get("tags", [])

    for section, original_key, translated_key in (
        ("executive_summary", "summary", "summary_translated"),
        ("meeting_brief", "brief", "brief_translated"),
    ):
        if sections.get(section):
            value, used_fallback = _translated(meeting.get(original_key), meeting.get(translated_key), language)
            document[section] = value
            if used_fallback:
                fallback.add(section)

    if sections.get("speakers"):
        document["speakers"] = data.get("speakers", [])

    if sections.get("transcript"):
        rows = []
        total_characters = 0
        for segment in sorted(data.get("transcript", []), key=lambda row: (row.get("segment_index", 0), row.get("start_ms", 0))):
            original = segment.get("original_text") or segment.get("text")
            text, used_fallback = _translated(original, segment.get("translated_text"), language)
            if used_fallback:
                fallback.add("transcript")
            total_characters += len(text or "")
            if total_characters > MAX_TRANSCRIPT_CHARACTERS:
                raise ValueError("Transcript is too large to export.")
            row = {"speaker": segment.get("display_name") or segment.get("speaker_label") or "Unknown Speaker", "text": text}
            if sections.get("timestamps"):
                row.update({"start_ms": segment.get("start_ms"), "end_ms": segment.get("end_ms")})
            rows.append(row)
        document["transcript"] = rows

    field_contracts = {
        "action_items": (("title", "translated_title"), ("description", "translated_description"), ("assignee", None), ("due_at", None), ("priority", None), ("status", None)),
        "decisions": (("title", "translated_title"), ("description", "translated_description")),
        "questions": (("question", "translated_question"), ("answer", "translated_answer"), ("status", None)),
    }
    for section, fields in field_contracts.items():
        if not sections.get(section):
            continue
        output_rows = []
        for source_row in data.get(section, []):
            output_row = {}
            for original_key, translated_key in fields:
                if translated_key:
                    value, used_fallback = _translated(source_row.get(original_key), source_row.get(translated_key), language)
                    if used_fallback:
                        fallback.add(section)
                else:
                    value = source_row.get(original_key)
                output_row[original_key] = value
            output_rows.append(output_row)
        document[section] = output_rows

    if sections.get("processing_timeline"):
        document["processing_timeline"] = data.get("processing_timeline", [])
    document["translation_fallback_sections"] = sorted(fallback)
    return document


SECTION_TITLES = {
    "meeting_details": "Meeting Details", "executive_summary": "Executive Summary",
    "meeting_brief": "Meeting Brief", "speakers": "Speakers", "transcript": "Transcript",
    "action_items": "Action Items", "decisions": "Decisions", "questions": "Questions",
    "tags": "Tags", "processing_timeline": "Processing Timeline",
}


def _plain_lines(document: dict[str, Any]) -> list[str]:
    lines = [document["title"], ""]
    fallbacks = set(document.get("translation_fallback_sections", []))
    for key, title in SECTION_TITLES.items():
        if key not in document:
            continue
        lines.extend([title, "=" * len(title)])
        if key in fallbacks:
            lines.append("Translation unavailable - original language shown.")
        value = document[key]
        if key == "transcript":
            for row in value:
                stamp = ""
                if "start_ms" in row:
                    stamp = f"[{format_timestamp(row['start_ms'])}-{format_timestamp(row['end_ms'])}] "
                lines.extend([f"{stamp}{row['speaker']}:", row.get("text") or "", ""])
        elif isinstance(value, dict):
            lines.extend(f"{name.replace('_', ' ').title()}: {item if item is not None else ''}" for name, item in value.items())
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    lines.append("; ".join(f"{name.replace('_', ' ').title()}: {field}" for name, field in item.items() if field not in (None, "")))
                else:
                    lines.append(str(item))
        elif value is not None:
            lines.append(str(value))
        lines.append("")
    return lines


def render_text(document: dict[str, Any]) -> bytes:
    return "\n".join(_plain_lines(document)).encode("utf-8")


def render_markdown(document: dict[str, Any]) -> bytes:
    lines = [f"# {document['title']}", ""]
    for line in _plain_lines(document)[2:]:
        if line and set(line) == {"="}:
            continue
        if line in SECTION_TITLES.values():
            lines.extend([f"## {line}", ""])
        else:
            lines.append(line)
    return "\n".join(lines).encode("utf-8")


def render_json(document: dict[str, Any]) -> bytes:
    return json.dumps(document, ensure_ascii=False, indent=2).encode("utf-8")


def render_docx(document: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        style = styles[name]; style.font.name = "Arial"; style.font.size = Pt(size); style.font.color.rgb = RGBColor(46, 116, 181)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(document["title"]); run.bold = True; run.font.name = "Arial"; run.font.size = Pt(24); run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph("MeetingVA Export")
    subtitle.paragraph_format.space_after = Pt(18)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("MeetingVA | ")
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field)
    for line in _plain_lines(document)[2:]:
        if line in SECTION_TITLES.values():
            doc.add_heading(line, level=1)
        elif line and set(line) == {"="}:
            continue
        else:
            doc.add_paragraph(line)
    output = BytesIO(); doc.save(output); return output.getvalue()


def render_pdf(document: dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    font_name = "Helvetica"
    for font_path in (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("C:/Windows/Fonts/arial.ttf")):
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("MeetingVAUnicode", str(font_path))); font_name = "MeetingVAUnicode"; break
    output = BytesIO()
    pdf = SimpleDocTemplate(output, pagesize=letter, leftMargin=0.8*inch, rightMargin=0.8*inch, topMargin=0.75*inch, bottomMargin=0.9*inch, title=document["title"])
    styles = getSampleStyleSheet()
    body = ParagraphStyle("MVBody", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=14, spaceAfter=7)
    heading = ParagraphStyle("MVHeading", parent=styles["Heading1"], fontName=font_name, textColor=colors.HexColor("#2E74B5"), fontSize=15, leading=18, spaceBefore=12, spaceAfter=7)
    title = ParagraphStyle("MVTitle", parent=styles["Title"], fontName=font_name, textColor=colors.HexColor("#0B2545"), fontSize=25, leading=30, alignment=TA_CENTER)
    story = [Spacer(1, 1.2*inch), Paragraph(_escape(document["title"]), title), Spacer(1, 0.2*inch), Paragraph("MeetingVA Export", body), PageBreak()]
    for line in _plain_lines(document)[2:]:
        if line in SECTION_TITLES.values(): story.append(Paragraph(_escape(line), heading))
        elif line and set(line) == {"="}: continue
        elif line: story.append(Paragraph(_escape(line), body))
        else: story.append(Spacer(1, 5))
    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState(); canvas.setFont(font_name, 8); canvas.setFillColor(colors.grey); canvas.drawRightString(letter[0]-0.8*inch, 0.4*inch, f"Page {doc.page}"); canvas.restoreState()
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()


def _escape(value: Any) -> str:
    from xml.sax.saxutils import escape
    return escape(str(value or "")).replace("\n", "<br/>")


def render_export(document: dict[str, Any], export_format: str) -> bytes:
    renderers = {"pdf": render_pdf, "docx": render_docx, "md": render_markdown, "txt": render_text, "json": render_json}
    payload = renderers[export_format](document)
    if len(payload) > MAX_EXPORT_BYTES:
        raise ValueError("Generated export is too large.")
    return payload
